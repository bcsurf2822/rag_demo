"""
Module for loading text and PDF files.
"""
import logging
import os
from typing import Dict, Any, Optional
from docx import Document

import PyPDF2

logger = logging.getLogger(__name__)

def load_text_file(file_path: str) -> str:
    """
    Load text from a text file.
    
    Args:
        file_path: Path to the text file.
        
    Returns:
        The text content of the file.
        
    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be read as text.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise ValueError(f"Could not read file {file_path}: {str(e)}")
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        raise ValueError(f"Could not read file {file_path}: {str(e)}")

def load_pdf_file(file_path: str) -> str:
    """
    Load text from a PDF file.
    
    Args:
        file_path: Path to the PDF file.
        
    Returns:
        The text content of the PDF.
        
    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be read as a PDF.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        return text
    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {str(e)}")
        raise ValueError(f"Could not read PDF {file_path}: {str(e)}")

def load_file(file_path: str) -> Dict[str, Any]:
    """
    Load text from a file based on its extension.
    
    Args:
        file_path: Path to the file.
        
    Returns:
        A dictionary containing the file metadata and content.
        
    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported or cannot be read.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    content = ""
    file_type = ""
    
    if file_ext == '.txt':
        content = load_text_file(file_path)
        file_type = "text"
    elif file_ext == '.pdf':
        content = load_pdf_file(file_path)
        file_type = "pdf"
    elif file_ext == '.docx':
        content = load_docx_file(file_path)
        file_type = "docx"
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    return {
        "filename": file_name,
        "title": os.path.splitext(file_name)[0],
        "content": content,
        "file_type": file_type
    }

def extract_file_content(file_obj: Any, file_name: Optional[str] = None) -> Dict[str, Any]:
    """
    Extract content from an uploaded file object (for Streamlit).
    
    Args:
        file_obj: The uploaded file object.
        file_name: Optional file name override.
        
    Returns:
        A dictionary containing the file metadata and content.
        
    Raises:
        ValueError: If the file type is not supported or cannot be read.
    """
    if file_name is None:
        file_name = file_obj.name
    
    file_ext = os.path.splitext(file_name)[1].lower()
    
    content = ""
    file_type = ""
    
    try:
        if file_ext == '.txt':
            content = file_obj.getvalue().decode('utf-8')
            file_type = "text"
        elif file_ext == '.pdf':
            pdf_reader = PyPDF2.PdfReader(file_obj)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                content += page.extract_text() + "\n\n"
            
            file_type = "pdf"
        elif file_ext == '.docx':
            doc = Document(file_obj)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + "\t"
                    content += "\n"
            
            file_type = "docx"
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        return {
            "filename": file_name,
            "title": os.path.splitext(file_name)[0],
            "content": content,
            "file_type": file_type
        }
    except Exception as e:
        logger.error(f"Error processing file {file_name}: {str(e)}")
        raise ValueError(f"Could not process file {file_name}: {str(e)}") 
    

def load_docx_file(file_path: str) -> str:
    """
    Load text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file.
        
    Returns:
        The text content of the DOCX file.
        
    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be read as a DOCX.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {str(e)}")
        raise ValueError(f"Could not read DOCX {file_path}: {str(e)}")